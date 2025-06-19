# generate_db_structure.py

import os
import django
from docx import Document
from django.apps import apps
from django.db import models

# Setup Django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "canteen.settings")  # Replace with your actual settings
django.setup()

doc = Document()
doc.add_heading("Database Table Relationships - Smart Canteen", level=1)

for model in apps.get_models():
    doc.add_heading(model.__name__, level=2)

    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field Name"
    hdr_cells[1].text = "Field Type"
    hdr_cells[2].text = "Related Model (if any)"

    for field in model._meta.get_fields():
        if isinstance(field, models.Field):
            row_cells = table.add_row().cells
            row_cells[0].text = field.name
            row_cells[1].text = field.get_internal_type()
            if isinstance(field, (models.ForeignKey, models.OneToOneField, models.ManyToManyField)):
                row_cells[2].text = f"{field.related_model.__name__}"
            else:
                row_cells[2].text = "-"

    doc.add_paragraph("")

doc.save("SmartCanteen_DB_Relations.docx")
print("✅ Document generated: SmartCanteen_DB_Relations.docx")
